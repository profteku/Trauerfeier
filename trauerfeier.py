import sys
import os
from docx import Document

def get_replacements(name, gender):
    """
    Gibt ein Dictionary mit den Platzhaltern und den
    entsprechenden Ersetzungen basierend auf dem Geschlecht zurück.
    """
    if gender.lower() == 'm':
        return {
            "NAME": name,
            "$P": "Er",
            "$R": "ihn",
            "$D": "ihm",
            "$V": "den Verstorbenen",
        }
    elif gender.lower() == 'w':
        return {
            "NAME": name,
            "$P": "Sie",
            "$R": "sie",
            "$D": "ihr",
            "$V": "die Verstorbene",
        }
    return {}

def update_word_document(name, gender):
    """
    Ersetzt die Platzhalter im Word-Dokument und speichert es im Ordner 'Output'.
    """
    try:
        doc = Document("Trauerfeier_Roh.docx")
        replacements = get_replacements(name, gender)

        def apply_replacements(element):
            for paragraph in element.paragraphs:
                for run in paragraph.runs:
                    for placeholder, value in replacements.items():
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)

        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        run.text = run.text.replace(placeholder, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    apply_replacements(cell)

        # --- NEUE ÄNDERUNGEN HIER ---
        # 1. Definiere den Namen des Output-Ordners
        output_folder = "Output"

        # 2. Erstelle den Ordner, falls er nicht existiert
        os.makedirs(output_folder, exist_ok=True)

        # 3. Erstelle den vollständigen Pfad für die neue Datei
        file_name = f"Trauerfeier_{name}.docx"
        output_path = os.path.join(output_folder, file_name)
        
        # 4. Speichere das Dokument im neuen Pfad
        doc.save(output_path)
        
        print(f"✅ Dokument erfolgreich in '{output_path}' erstellt.")
        # --- ENDE DER ÄNDERUNGEN ---

    except FileNotFoundError:
        print("❌ Fehler: Die Datei 'Trauerfeier_Roh.docx' wurde nicht gefunden.")
    except Exception as e:
        print(f"Ein unerwarteter Fehler ist aufgetreten: {e}")

if __name__ == "__main__":
    if len(sys.argv) > 2:
        name_to_insert = sys.argv[1]
        gender_arg = sys.argv[2].lower()
        
        if gender_arg not in ['m', 'w']:
            print("❌ Fehler: Das zweite Argument für das Geschlecht muss 'm' oder 'w' sein.")
        else:
            update_word_document(name_to_insert, gender_arg)
    else:
        print("ℹ️ Bitte gib einen Namen und das Geschlecht ('m' oder 'w') an.")
        print("   Beispiel (männlich): python trauerfeier_generator.py 'Max Mustermann' m")
        print("   Beispiel (weiblich): python trauerfeier_generator.py 'Erika Mustermann' w")